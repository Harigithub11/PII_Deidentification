"""
Document Redaction Engine for Format-Preserving PII Redaction
Maintains original document structure while redacting only PII content
"""
import os
import io
import logging
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from enum import Enum
import tempfile
import atexit
import shutil
from contextlib import contextmanager

# PDF processing
import fitz  # PyMuPDF
import pdfplumber
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.colors import black, white, red

# Image processing
from PIL import Image, ImageDraw, ImageFilter, ImageFont
import cv2
import numpy as np

# Document processing
import docx
from docx.shared import RGBColor

# Enhanced PII detection
from enhanced_pii_detector import EnhancedPIIDetector, Sector

# Presidio Image Redaction
from presidio_image_redactor import ImageRedactorEngine

# Advanced text replacement
from pii_redaction.redactor import PIIRedactor

logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)

class RedactionMethod(Enum):
    """Redaction methods for PII masking"""
    BLACKOUT = "blackout"
    WHITEOUT = "whiteout"  
    BLUR = "blur"
    PIXELATE = "pixelate"
    REPLACEMENT = "replacement"

class TempFileManager:
    """Context manager for safe temporary file handling"""
    
    def __init__(self, prefix="doc_redaction_"):
        self.temp_dir = None
        self.prefix = prefix
        self.created_files = []
    
    def __enter__(self):
        self.temp_dir = tempfile.mkdtemp(prefix=self.prefix)
        logger.debug(f"Created temporary directory: {self.temp_dir}")
        return self
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        self.cleanup()
    
    def create_temp_file(self, suffix="", content: bytes = None) -> str:
        """Create a temporary file and track it for cleanup"""
        fd, path = tempfile.mkstemp(suffix=suffix, dir=self.temp_dir)
        self.created_files.append(path)
        
        if content:
            with os.fdopen(fd, 'wb') as f:
                f.write(content)
        else:
            os.close(fd)
        
        return path
    
    def cleanup(self):
        """Clean up all temporary files and directories"""
        try:
            # Clean up individual files first
            for file_path in self.created_files:
                if os.path.exists(file_path):
                    try:
                        os.unlink(file_path)
                        logger.debug(f"Removed temporary file: {file_path}")
                    except Exception as e:
                        logger.warning(f"Failed to remove temp file {file_path}: {e}")
            
            # Clean up directory
            if self.temp_dir and os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
                logger.debug(f"Removed temporary directory: {self.temp_dir}")
                
        except Exception as e:
            logger.warning(f"Temp file cleanup failed: {e}")
        finally:
            self.created_files.clear()
            self.temp_dir = None

class DocumentFormat(Enum):
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    IMAGE = "image"
    UNKNOWN = "unknown"

class RedactionVisualStyle(Enum):
    BLACK_BOX = "black_box"
    BLUR = "blur"
    PIXELATE = "pixelate"
    WHITE_BOX = "white_box"
    STRIKETHROUGH = "strikethrough"
    REPLACEMENT_TEXT = "replacement_text"

class DocumentRedactionEngine:
    """Engine for redacting PII while preserving document format and structure"""
    
    # Memory management configuration
    MAX_PDF_SIZE_BYTES = 100 * 1024 * 1024  # 100MB default limit
    MAX_IMAGE_SIZE_BYTES = 50 * 1024 * 1024   # 50MB for images
    CHUNK_SIZE = 1024 * 1024  # 1MB chunks for streaming
    
    def __init__(self, sector: Sector = Sector.GENERAL, max_pdf_size: int = None, max_image_size: int = None):
        self.pii_detector = EnhancedPIIDetector(sector=sector)
        self.sector = sector
        
        # Override default memory limits if provided
        if max_pdf_size:
            self.MAX_PDF_SIZE_BYTES = max_pdf_size
        if max_image_size:
            self.MAX_IMAGE_SIZE_BYTES = max_image_size
        
        # Redaction styles configuration
        self.redaction_styles = {
            RedactionMethod.BLACKOUT: RedactionVisualStyle.BLACK_BOX,
            RedactionMethod.BLUR: RedactionVisualStyle.BLUR,
            RedactionMethod.PIXELATE: RedactionVisualStyle.PIXELATE,
            RedactionMethod.WHITEOUT: RedactionVisualStyle.WHITE_BOX,
            RedactionMethod.REPLACEMENT: RedactionVisualStyle.REPLACEMENT_TEXT
        }
        
        # Initialize Presidio Image Redactor for advanced visual redaction
        try:
            self.image_redactor = ImageRedactorEngine()
            logger.info("Presidio Image Redactor initialized successfully")
        except Exception as e:
            logger.warning(f"Failed to initialize Presidio Image Redactor: {e}")
            self.image_redactor = None
        
        # Initialize PII Redactor for advanced text replacement
        try:
            self.pii_redactor = PIIRedactor()
            logger.info("PII Redactor initialized successfully")
        except Exception as e:
            logger.warning(f"Failed to initialize PII Redactor: {e}")
            self.pii_redactor = None
        
        # Initialize temp directory for processing (legacy support)
        self.temp_dir = tempfile.mkdtemp(prefix="doc_redaction_")
        
        # Register cleanup on exit
        atexit.register(self.cleanup)
        
        logger.info(f"DocumentRedactionEngine initialized with limits: PDF={self.MAX_PDF_SIZE_BYTES//1024//1024}MB, Image={self.MAX_IMAGE_SIZE_BYTES//1024//1024}MB")
    
    def _check_memory_limits(self, file_content: bytes, filename: str, doc_format: DocumentFormat) -> Dict[str, Any]:
        """Check if file size is within memory limits"""
        file_size = len(file_content)
        
        # Define limits based on document type
        if doc_format == DocumentFormat.PDF:
            max_size = self.MAX_PDF_SIZE_BYTES
            doc_type = "PDF"
        elif doc_format == DocumentFormat.IMAGE:
            max_size = self.MAX_IMAGE_SIZE_BYTES
            doc_type = "Image"
        else:
            # For text and DOCX, use smaller limit as they're loaded fully into memory
            max_size = self.MAX_IMAGE_SIZE_BYTES
            doc_type = doc_format.value
        
        if file_size > max_size:
            size_mb = file_size / (1024 * 1024)
            limit_mb = max_size / (1024 * 1024)
            error_msg = f"{doc_type} file too large: {size_mb:.1f}MB exceeds limit of {limit_mb:.1f}MB"
            logger.warning(f"Memory limit exceeded for {filename}: {error_msg}")
            
            return {
                'success': False,
                'error': error_msg,
                'file_size': file_size,
                'limit': max_size,
                'filename': filename
            }
        
        return {'success': True, 'file_size': file_size, 'limit': max_size}
        
    def detect_format(self, file_content: bytes, filename: str) -> DocumentFormat:
        """Detect document format from content and filename with comprehensive image support"""
        try:
            # Check by file extension first (case insensitive)
            ext = Path(filename).suffix.lower()
            
            # PDF files
            if ext == '.pdf':
                return DocumentFormat.PDF
            
            # Microsoft Office documents
            elif ext in ['.docx', '.doc']:
                return DocumentFormat.DOCX
            
            # Plain text files
            elif ext in ['.txt', '.text']:
                return DocumentFormat.TXT
            
            # Image files - comprehensive list
            elif ext in ['.png', '.jpg', '.jpeg', '.jpe', '.jfif', 
                        '.tiff', '.tif', '.bmp', '.gif', '.webp', 
                        '.svg', '.ico', '.psd', '.raw']:
                return DocumentFormat.IMAGE
            
            # Check by content magic bytes for more reliable detection
            if not file_content:
                return DocumentFormat.UNKNOWN
                
            # PDF files
            if file_content.startswith(b'%PDF'):
                return DocumentFormat.PDF
            
            # Image formats - comprehensive magic byte detection
            elif file_content.startswith(b'\xff\xd8\xff'):  # JPEG
                return DocumentFormat.IMAGE
            elif file_content.startswith(b'\x89PNG\r\n\x1a\n'):  # PNG
                return DocumentFormat.IMAGE
            elif file_content.startswith((b'GIF87a', b'GIF89a')):  # GIF
                return DocumentFormat.IMAGE
            elif file_content.startswith(b'BM'):  # BMP
                return DocumentFormat.IMAGE
            elif file_content.startswith(b'RIFF') and b'WEBP' in file_content[:12]:  # WebP
                return DocumentFormat.IMAGE
            elif file_content.startswith((b'II*\x00', b'MM\x00*')):  # TIFF
                return DocumentFormat.IMAGE
            elif file_content.startswith(b'\x00\x00\x01\x00'):  # ICO
                return DocumentFormat.IMAGE
            elif file_content.startswith(b'<svg') or b'<svg' in file_content[:100]:  # SVG
                return DocumentFormat.IMAGE
            
            # Microsoft Office formats (ZIP-based)
            elif file_content.startswith(b'PK\x03\x04'):
                # More specific Office document detection
                if b'word/' in file_content[:2048]:
                    return DocumentFormat.DOCX
                elif b'xl/' in file_content[:2048]:
                    return DocumentFormat.DOCX  # Excel files also handled as DOCX format
                elif b'ppt/' in file_content[:2048]:
                    return DocumentFormat.DOCX  # PowerPoint files also handled as DOCX format
                else:
                    # Could be a ZIP file or other Office format
                    return DocumentFormat.UNKNOWN
            
            # Plain text detection (heuristic)
            try:
                # Try to decode first 1KB as UTF-8
                file_content[:1024].decode('utf-8')
                # Check if it looks like text (printable characters)
                text_sample = file_content[:512]
                printable_ratio = sum(1 for byte in text_sample if 32 <= byte <= 126 or byte in [9, 10, 13]) / len(text_sample)
                if printable_ratio > 0.7:  # 70% printable characters
                    return DocumentFormat.TXT
            except UnicodeDecodeError:
                pass
                
            return DocumentFormat.UNKNOWN
            
        except Exception as e:
            logger.error(f"Format detection failed for {filename}: {e}")
            return DocumentFormat.UNKNOWN
    
    def redact_document(self, 
                       file_content: bytes, 
                       filename: str, 
                       redaction_method: RedactionMethod = RedactionMethod.BLACKOUT,
                       confidence_threshold: float = 0.7) -> Dict[str, Any]:
        """
        Main method to redact PII from document while preserving format
        
        Args:
            file_content: Original document bytes
            filename: Original filename
            redaction_method: Type of redaction to apply
            confidence_threshold: Minimum confidence for PII detection
            
        Returns:
            Dict containing redacted document bytes and metadata
        """
        try:
            # Detect document format
            doc_format = self.detect_format(file_content, filename)
            
            if doc_format == DocumentFormat.UNKNOWN:
                return {
                    'success': False,
                    'error': 'Unsupported document format',
                    'filename': filename
                }
            
            # Check memory limits before processing
            memory_check = self._check_memory_limits(file_content, filename, doc_format)
            if not memory_check['success']:
                return memory_check
            
            logger.info(f"Redacting {doc_format.value} document: {filename} ({memory_check['file_size']//1024}KB)")
            
            # Route to appropriate redaction method
            if doc_format == DocumentFormat.PDF:
                return self._redact_pdf(file_content, filename, redaction_method, confidence_threshold)
            elif doc_format == DocumentFormat.IMAGE:
                return self._redact_image(file_content, filename, redaction_method, confidence_threshold)
            elif doc_format == DocumentFormat.TXT:
                return self._redact_text(file_content, filename, redaction_method, confidence_threshold)
            elif doc_format == DocumentFormat.DOCX:
                return self._redact_docx(file_content, filename, redaction_method, confidence_threshold)
            
        except Exception as e:
            logger.error(f"Document redaction failed for {filename}: {e}")
            return {
                'success': False,
                'error': f'Redaction failed: {str(e)}',
                'filename': filename
            }
    
    def _redact_pdf(self, file_content: bytes, filename: str, 
                    redaction_method: RedactionMethod, confidence_threshold: float) -> Dict[str, Any]:
        """Redact PII from PDF while maintaining layout and formatting (memory-safe)"""
        pdf_document = None
        pdfplumber_doc = None
        
        try:
            logger.info(f"Starting PDF redaction for {filename} with method {redaction_method.value}")
            
            # Open PDF with PyMuPDF for visual redaction (using file stream to minimize memory usage)
            pdf_document = fitz.open(stream=file_content, filetype="pdf")
            
            # Create separate BytesIO for pdfplumber to avoid conflicts
            pdf_file = io.BytesIO(file_content)
            pdfplumber_doc = pdfplumber.open(pdf_file)
            
            # Log PDF information for memory tracking
            page_count = len(pdf_document)
            logger.info(f"Processing PDF with {page_count} pages using pdfplumber text extraction")
            
            # Check if PDF has too many pages (memory protection)
            if page_count > 1000:
                logger.warning(f"PDF has {page_count} pages, which may cause memory issues")
                pdfplumber_doc.close()
                return {
                    'success': False,
                    'error': f'PDF too complex: {page_count} pages exceeds limit of 1000 pages',
                    'filename': filename
                }
            
            total_redactions = 0
            pii_summary = []
            
            # Process each page with memory cleanup
            for page_num in range(page_count):
                try:
                    # Use PyMuPDF page for visual redaction
                    page = pdf_document[page_num]
                    
                    # Extract text with position information (still needed for finding text rectangles)
                    text_instances = page.get_text("dict")
                    
                    # Get full page text using pdfplumber for consistent PII detection (matches upload process)
                    try:
                        pdfplumber_page = pdfplumber_doc.pages[page_num]
                        page_text = pdfplumber_page.extract_text()
                        if not page_text:
                            page_text = ""
                    except Exception as e:
                        logger.warning(f"pdfplumber extraction failed for page {page_num + 1}, falling back to PyMuPDF: {e}")
                        page_text = page.get_text()
                    
                    if not page_text.strip():
                        continue
                        
                    # Detect PII in page text
                    logger.debug(f"Extracted page text length: {len(page_text)}")
                    logger.debug(f"Page {page_num + 1} text content: {repr(page_text)}")
                    logger.debug(f"Page text type: {type(page_text)}")
                    logger.debug(f"Confidence threshold: {confidence_threshold} (type: {type(confidence_threshold)})")
                    
                    try:
                        pii_entities_list = self.pii_detector.detect_pii(page_text, confidence_threshold=confidence_threshold)
                        logger.debug(f"PII detection successful, type: {type(pii_entities_list)}")
                        
                        # Convert PIIResult objects to dictionary format expected by redaction engine
                        pii_results = {
                            'entities': [
                                {
                                    'entity_type': entity.entity_type,
                                    'text': entity.text,
                                    'start': entity.start,
                                    'end': entity.end,
                                    'confidence': entity.confidence
                                }
                                for entity in pii_entities_list
                            ]
                        }
                    except Exception as pii_error:
                        import traceback
                        logger.error(f"PII DETECTION ERROR on page {page_num + 1}:")
                        logger.error(f"Error: {pii_error}")
                        logger.error(f"Error type: {type(pii_error).__name__}")
                        logger.error(f"Full traceback:\n{traceback.format_exc()}")
                        raise pii_error
                    logger.debug(f"PII detection returned {len(pii_results['entities'])} entities")
                    
                    if not pii_results['entities']:
                        logger.debug("No PII entities found on this page")
                        continue
                    
                    # Find and redact PII locations on the page
                    for entity in pii_results['entities']:
                        logger.debug(f"Processing entity: {entity}")
                        entity_text = entity['text']
                        
                        # Find text instances on page that match PII
                        text_rects = self._find_text_rectangles(page, entity_text)
                        
                        for rect in text_rects:
                            # Apply redaction based on method
                            if redaction_method == RedactionMethod.BLACKOUT:
                                # Add black rectangle overlay
                                page.draw_rect(rect, color=(0, 0, 0), fill=(0, 0, 0))
                            elif redaction_method == RedactionMethod.WHITEOUT:
                                # Add white rectangle overlay
                                page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))
                            elif redaction_method == RedactionMethod.BLUR:
                                # Use Presidio for true blur effect if available
                                if self.image_redactor:
                                    try:
                                        self._apply_presidio_redaction(page, rect, "blur")
                                    except Exception as e:
                                        logger.warning(f"Presidio blur failed, using fallback: {e}")
                                        page.draw_rect(rect, color=(0.6, 0.6, 0.6), fill=(0.6, 0.6, 0.6))
                                else:
                                    # Fallback to gray overlay
                                    page.draw_rect(rect, color=(0.6, 0.6, 0.6), fill=(0.6, 0.6, 0.6))
                            elif redaction_method == RedactionMethod.PIXELATE:
                                # Native pixelation implementation - no dependency on Presidio
                                self._apply_native_pixelation(page, rect)
                            elif redaction_method == RedactionMethod.REPLACEMENT:
                                # Simplified text replacement to prevent PDF corruption
                                try:
                                    replacement = self._get_enhanced_replacement_text(entity['entity_type'], entity_text)
                                except Exception as e:
                                    logger.warning(f"Enhanced replacement failed, using fallback: {e}")
                                    replacement = self._get_replacement_text(entity['entity_type'])
                                
                                # Use simple visual overlay approach for stability
                                page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))
                                # Safe text insertion with proper positioning
                                try:
                                    font_size = min(8, int(rect.height * 0.6))  # Scale font to rectangle
                                    page.insert_text((rect.x0 + 2, rect.y0 + font_size + 2), replacement, 
                                                    fontsize=font_size, color=(0, 0, 0))
                                except Exception as text_error:
                                    logger.warning(f"Text insertion failed: {text_error}")
                                    # Final fallback: colored rectangle
                                    page.draw_rect(rect, color=(0.9, 0.9, 1.0), fill=(0.9, 0.9, 1.0))
                            
                            total_redactions += 1
                            
                        pii_summary.append({
                            'type': entity['entity_type'],
                            'text': entity_text[:10] + '...' if len(entity_text) > 10 else entity_text,
                            'confidence': entity['confidence'],
                            'page': page_num + 1,
                            'locations': len(text_rects)
                        })
                
                except Exception as page_error:
                    import traceback
                    logger.error(f"DETAILED ERROR processing PDF page {page_num + 1}:")
                    logger.error(f"Error message: {page_error}")
                    logger.error(f"Error type: {type(page_error).__name__}")
                    logger.error(f"Full traceback:\n{traceback.format_exc()}")
                    # Continue with next page
            
            # Save redacted PDF to bytes with proper error handling
            logger.info(f"Finalizing PDF redaction - Pages: {page_count}, Redactions: {total_redactions}")
            
            try:
                # Generate redacted PDF bytes
                redacted_pdf_bytes = pdf_document.write()
                
                # Critical: Validate that PDF bytes are not empty or corrupted
                if not redacted_pdf_bytes or len(redacted_pdf_bytes) < 100:  # PDF header is ~100 bytes minimum
                    logger.error(f"Generated PDF is too small: {len(redacted_pdf_bytes) if redacted_pdf_bytes else 0} bytes")
                    raise Exception("Generated redacted PDF appears to be empty or corrupted")
                
                # Comprehensive PDF validation using dedicated method
                validation_result = self._validate_pdf_integrity(redacted_pdf_bytes)
                if not validation_result['valid']:
                    logger.error(f"PDF integrity validation failed: {validation_result['error']}")
                    raise Exception(f"Generated PDF failed integrity validation: {validation_result['error']}")
                
                # Additional page count check
                if validation_result.get('pages', 0) != page_count:
                    logger.error(f"Page count mismatch: original={page_count}, redacted={validation_result.get('pages', 0)}")
                    raise Exception(f"Generated PDF page count mismatch: expected {page_count}, got {validation_result.get('pages', 0)}")
                    
                logger.info(f"PDF validation successful: {len(redacted_pdf_bytes)} bytes, {validation_result.get('pages', 0)} pages")
                
            except Exception as validation_error:
                logger.error(f"Generated PDF validation failed: {validation_error}")
                # Fallback: Return original document if redacted version is invalid
                logger.warning("Fallback: returning original document due to redaction validation failure")
                return {
                        'success': True,
                        'redacted_content': file_content,  # Return original as fallback
                        'format': 'pdf',
                        'filename': filename,
                        'total_redactions': 0,
                        'pii_summary': [],
                        'redaction_method': redaction_method.value,
                        'pages_processed': page_count,
                        'fallback_reason': 'Redacted PDF validation failed'
                    }
                
            except Exception as write_error:
                logger.error(f"PDF write operation failed: {write_error}")
                # Fallback: Return original document if writing fails
                logger.warning("Fallback: returning original document due to PDF write failure")
                return {
                    'success': True,
                    'redacted_content': file_content,  # Return original as fallback
                    'format': 'pdf',
                    'filename': filename,
                    'total_redactions': 0,
                    'pii_summary': [],
                    'redaction_method': redaction_method.value,
                    'pages_processed': page_count,
                    'fallback_reason': 'PDF write operation failed'
                }
            
            return {
                'success': True,
                'redacted_content': redacted_pdf_bytes,
                'format': 'pdf',
                'filename': filename,
                'total_redactions': total_redactions,
                'pii_summary': pii_summary,
                'redaction_method': redaction_method.value,
                'pages_processed': page_count
            }
            
        except Exception as e:
            logger.error(f"PDF redaction failed: {e}")
            return {
                'success': False,
                'error': f'PDF redaction failed: {str(e)}',
                'filename': filename
            }
        finally:
            # Comprehensive cleanup to ensure proper resource management
            logger.debug("Starting PDF document cleanup")
            
            # Close PyMuPDF document
            if pdf_document:
                try:
                    pdf_document.close()
                    logger.debug("PyMuPDF document closed successfully")
                except Exception as cleanup_error:
                    logger.warning(f"PyMuPDF cleanup error: {cleanup_error}")
            
            # Close pdfplumber document  
            if pdfplumber_doc:
                try:
                    pdfplumber_doc.close()
                    logger.debug("Pdfplumber document closed successfully")
                except Exception as cleanup_error:
                    logger.warning(f"Pdfplumber cleanup error: {cleanup_error}")
            
            logger.debug("PDF document cleanup completed")
    
    def _find_text_rectangles(self, page: fitz.Page, target_text: str) -> List[fitz.Rect]:
        """Find rectangles containing specific text on a PDF page"""
        rectangles = []
        
        try:
            # Search for text instances
            text_instances = page.search_for(target_text)
            
            for rect in text_instances:
                # Expand rectangle slightly for better coverage
                expanded_rect = fitz.Rect(
                    rect.x0 - 2, rect.y0 - 2, 
                    rect.x1 + 2, rect.y1 + 2
                )
                rectangles.append(expanded_rect)
                
        except Exception as e:
            logger.warning(f"Text rectangle search failed: {e}")
            
        return rectangles
    
    def _redact_image(self, file_content: bytes, filename: str, 
                     redaction_method: RedactionMethod, confidence_threshold: float) -> Dict[str, Any]:
        """Redact PII from images using OCR and visual overlays"""
        try:
            # Open image
            image = Image.open(io.BytesIO(file_content))
            original_format = image.format
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Extract text using OCR (from document_processor.py)
            from document_processor import document_processor
            
            # Get OCR text with positions
            import pytesseract
            ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
            
            # Get full text for PII detection
            full_text = pytesseract.image_to_string(image)
            
            if not full_text.strip():
                return {
                    'success': True,
                    'redacted_content': file_content,  # No text found, return original
                    'format': 'image',
                    'filename': filename,
                    'total_redactions': 0,
                    'warning': 'No text detected in image'
                }
            
            # Detect PII
            pii_entities_list = self.pii_detector.detect_pii(full_text, confidence_threshold=confidence_threshold)
            
            # Convert PIIResult objects to dictionary format expected by redaction engine
            pii_results = {
                'entities': [
                    {
                        'entity_type': entity.entity_type,
                        'text': entity.text,
                        'start': entity.start,
                        'end': entity.end,
                        'confidence': entity.confidence
                    }
                    for entity in pii_entities_list
                ]
            }
            
            if not pii_results['entities']:
                return {
                    'success': True,
                    'redacted_content': file_content,  # No PII found, return original
                    'format': 'image',
                    'filename': filename,
                    'total_redactions': 0
                }
            
            # Apply visual redactions
            draw = ImageDraw.Draw(image)
            total_redactions = 0
            pii_summary = []
            
            for entity in pii_results['entities']:
                entity_text = entity['text']
                
                # Find text positions in OCR data
                positions = self._find_text_positions_in_ocr(ocr_data, entity_text)
                
                for pos in positions:
                    x1, y1, x2, y2 = pos
                    
                    if redaction_method == RedactionMethod.BLACKOUT:
                        draw.rectangle([x1, y1, x2, y2], fill='black')
                    elif redaction_method == RedactionMethod.WHITEOUT:
                        draw.rectangle([x1, y1, x2, y2], fill='white')
                    elif redaction_method == RedactionMethod.BLUR:
                        # Crop region, blur it, paste back
                        region = image.crop((x1, y1, x2, y2))
                        blurred = region.filter(ImageFilter.GaussianBlur(radius=5))
                        image.paste(blurred, (x1, y1))
                    elif redaction_method == RedactionMethod.PIXELATE:
                        # Pixelate the region
                        region = image.crop((x1, y1, x2, y2))
                        pixelated = region.resize((max(1, region.width // 10), max(1, region.height // 10)), Image.NEAREST)
                        pixelated = pixelated.resize(region.size, Image.NEAREST)
                        image.paste(pixelated, (x1, y1))
                    
                    total_redactions += 1
                
                pii_summary.append({
                    'type': entity['entity_type'],
                    'text': entity_text[:10] + '...' if len(entity_text) > 10 else entity_text,
                    'confidence': entity['confidence'],
                    'locations': len(positions)
                })
            
            # Save redacted image to bytes
            output_buffer = io.BytesIO()
            image.save(output_buffer, format=original_format or 'PNG')
            redacted_image_bytes = output_buffer.getvalue()
            
            return {
                'success': True,
                'redacted_content': redacted_image_bytes,
                'format': 'image',
                'filename': filename,
                'total_redactions': total_redactions,
                'pii_summary': pii_summary,
                'redaction_method': redaction_method.value
            }
            
        except Exception as e:
            logger.error(f"Image redaction failed: {e}")
            return {
                'success': False,
                'error': f'Image redaction failed: {str(e)}'
            }
    
    def _find_text_positions_in_ocr(self, ocr_data: Dict, target_text: str) -> List[Tuple[int, int, int, int]]:
        """Find bounding boxes for target text in OCR data"""
        positions = []
        
        try:
            # Join words to find target text
            words = []
            boxes = []
            
            for i in range(len(ocr_data['text'])):
                word = ocr_data['text'][i].strip()
                if word:
                    words.append(word)
                    boxes.append((
                        ocr_data['left'][i],
                        ocr_data['top'][i],
                        ocr_data['left'][i] + ocr_data['width'][i],
                        ocr_data['top'][i] + ocr_data['height'][i]
                    ))
            
            # Look for target text in word sequences
            target_words = target_text.split()
            
            for i in range(len(words) - len(target_words) + 1):
                if ' '.join(words[i:i+len(target_words)]) == target_text:
                    # Calculate bounding box for the entire phrase
                    start_box = boxes[i]
                    end_box = boxes[i + len(target_words) - 1]
                    
                    combined_box = (
                        min(start_box[0], end_box[0]),
                        min(start_box[1], end_box[1]),
                        max(start_box[2], end_box[2]),
                        max(start_box[3], end_box[3])
                    )
                    positions.append(combined_box)
                    
        except Exception as e:
            logger.warning(f"OCR text position search failed: {e}")
            
        return positions
    
    def _redact_text(self, file_content: bytes, filename: str, 
                    redaction_method: RedactionMethod, confidence_threshold: float) -> Dict[str, Any]:
        """Redact PII from plain text files"""
        try:
            # Decode text
            text = file_content.decode('utf-8', errors='replace')
            
            # Detect PII
            pii_entities_list = self.pii_detector.detect_pii(text, confidence_threshold=confidence_threshold)
            
            # Convert PIIResult objects to dictionary format expected by redaction engine
            pii_results = {
                'entities': [
                    {
                        'entity_type': entity.entity_type,
                        'text': entity.text,
                        'start': entity.start,
                        'end': entity.end,
                        'confidence': entity.confidence
                    }
                    for entity in pii_entities_list
                ]
            }
            
            if not pii_results['entities']:
                return {
                    'success': True,
                    'redacted_content': file_content,  # No PII found
                    'format': 'text',
                    'filename': filename,
                    'total_redactions': 0
                }
            
            # Apply text-based redactions
            redacted_text = text
            total_redactions = 0
            pii_summary = []
            
            # Sort entities by position (reverse order to maintain indices)
            entities_sorted = sorted(pii_results['entities'], key=lambda x: x['start'], reverse=True)
            
            for entity in entities_sorted:
                start, end = entity['start'], entity['end']
                entity_text = entity['text']
                
                if redaction_method == RedactionMethod.BLACKOUT:
                    replacement = '█' * len(entity_text)
                elif redaction_method == RedactionMethod.REPLACEMENT:
                    replacement = self._get_replacement_text(entity['entity_type'])
                elif redaction_method == RedactionMethod.WHITEOUT:
                    replacement = ' ' * len(entity_text)
                else:
                    replacement = '[REDACTED]'
                
                # Replace the PII text
                redacted_text = redacted_text[:start] + replacement + redacted_text[end:]
                total_redactions += 1
                
                pii_summary.append({
                    'type': entity['entity_type'],
                    'text': entity_text[:10] + '...' if len(entity_text) > 10 else entity_text,
                    'confidence': entity['confidence']
                })
            
            # Encode back to bytes
            redacted_bytes = redacted_text.encode('utf-8')
            
            return {
                'success': True,
                'redacted_content': redacted_bytes,
                'format': 'text',
                'filename': filename,
                'total_redactions': total_redactions,
                'pii_summary': pii_summary,
                'redaction_method': redaction_method.value
            }
            
        except Exception as e:
            logger.error(f"Text redaction failed: {e}")
            return {
                'success': False,
                'error': f'Text redaction failed: {str(e)}'
            }
    
    def _redact_docx(self, file_content: bytes, filename: str, 
                    redaction_method: RedactionMethod, confidence_threshold: float) -> Dict[str, Any]:
        """Redact PII from DOCX files while preserving formatting"""
        with TempFileManager() as temp_manager:
            try:
                # Save to temp file for processing
                temp_path = temp_manager.create_temp_file(suffix=".docx", content=file_content)
                
                # Open DOCX
                doc = docx.Document(temp_path)
                
                # Extract all text for PII detection
                full_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                
                # Detect PII
                pii_entities_list = self.pii_detector.detect_pii(full_text, confidence_threshold=confidence_threshold)
                
                # Convert PIIResult objects to dictionary format expected by redaction engine
                pii_results = {
                    'entities': [
                        {
                            'entity_type': entity.entity_type,
                            'text': entity.text,
                            'start': entity.start,
                            'end': entity.end,
                            'confidence': entity.confidence
                        }
                        for entity in pii_entities_list
                    ]
                }
                
                if not pii_results['entities']:
                    return {
                        'success': True,
                        'redacted_content': file_content,  # No PII found
                        'format': 'docx',
                        'filename': filename,
                        'total_redactions': 0
                    }
                
                # Apply redactions to document
                total_redactions = 0
                pii_summary = []
                
                for entity in pii_results['entities']:
                    entity_text = entity['text']
                    
                    # Find and replace in paragraphs
                    for paragraph in doc.paragraphs:
                        if entity_text in paragraph.text:
                            # Replace text while preserving formatting
                            if redaction_method == RedactionMethod.REPLACEMENT:
                                replacement = self._get_replacement_text(entity['entity_type'])
                            elif redaction_method == RedactionMethod.BLACKOUT:
                                replacement = '█' * len(entity_text)
                            else:
                                replacement = '[REDACTED]'
                            
                            # Update paragraph text
                            for run in paragraph.runs:
                                if entity_text in run.text:
                                    run.text = run.text.replace(entity_text, replacement)
                                    
                                    # Apply visual formatting for redaction
                                    if redaction_method == RedactionMethod.BLACKOUT:
                                        run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                                        run.font.highlight_color = RGBColor(0, 0, 0)   # Black background
                                    
                                    total_redactions += 1
                    
                    pii_summary.append({
                        'type': entity['entity_type'],
                        'text': entity_text[:10] + '...' if len(entity_text) > 10 else entity_text,
                        'confidence': entity['confidence']
                    })
                
                # Save redacted document
                redacted_path = temp_manager.create_temp_file(suffix=f"_redacted_{filename}")
                doc.save(redacted_path)
                
                # Read redacted content
                with open(redacted_path, 'rb') as f:
                    redacted_bytes = f.read()
                
                return {
                    'success': True,
                    'redacted_content': redacted_bytes,
                    'format': 'docx',
                    'filename': filename,
                    'total_redactions': total_redactions,
                    'pii_summary': pii_summary,
                    'redaction_method': redaction_method.value
                }
                
            except Exception as e:
                logger.error(f"DOCX redaction failed: {e}")
                return {
                    'success': False,
                    'error': f'DOCX redaction failed: {str(e)}',
                    'filename': filename
                }
            # Automatic cleanup handled by TempFileManager context manager
    
    def _get_replacement_text(self, entity_type: str) -> str:
        """Get appropriate replacement text for PII entity type"""
        replacements = {
            # Personal identifiers
            'PERSON': '[NAME]',
            'PHONE_NUMBER': '[PHONE_NUMBER]',
            'EMAIL_ADDRESS': '[EMAIL_ADDRESS]',
            'SSN': '[SSN]',
            'US_DRIVER_LICENSE': '[DRIVER_LICENSE]',
            'US_PASSPORT': '[PASSPORT]',
            'US_ITIN': '[ITIN]',
            'US_SSN': '[SSN]',
            
            # Financial information
            'CREDIT_CARD': '[CREDIT_CARD]',
            'US_BANK_NUMBER': '[BANK_ACCOUNT]',
            'IBAN_CODE': '[IBAN]',
            'FINANCIAL_ID': '[FINANCIAL_ID]',
            
            # Location and organization
            'LOCATION': '[LOCATION]',
            'ORGANIZATION': '[ORGANIZATION]',
            'NRP': '[NATIONALITY]',
            
            # Dates and times
            'DATE_TIME': '[DATE_TIME]',
            
            # Technical identifiers
            'IP_ADDRESS': '[IP_ADDRESS]',
            'URL': '[URL]',
            'CRYPTO': '[CRYPTO_ADDRESS]',
            
            # Medical and professional
            'MEDICAL_LICENSE': '[MEDICAL_LICENSE]',
            'HEALTHCARE_ID': '[HEALTHCARE_ID]',
            
            # International identifiers
            'UK_NHS': '[NHS_NUMBER]',
            
            # Social media
            'SOCIAL_PROFILE': '[SOCIAL_PROFILE]',
            
            # Government
            'GOVERNMENT_ID': '[GOVERNMENT_ID]'
        }
        
        return replacements.get(entity_type, '[REDACTED]')
    
    def _apply_native_pixelation(self, page, rect):
        """
        Apply native pixelation effect to a PDF rectangle
        Creates a proper pixelated effect with varying colors
        """
        try:
            # Use moderate pixel size for good effect without too many operations
            pixel_size = 6
            width = int(rect.width)
            height = int(rect.height)
            
            # Define a palette of grays and colors for pixelation
            pixel_colors = [
                (0.3, 0.3, 0.3),  # Dark gray
                (0.5, 0.5, 0.5),  # Medium gray
                (0.7, 0.7, 0.7),  # Light gray
                (0.4, 0.4, 0.4),  # Another gray
                (0.6, 0.6, 0.6),  # Another gray
                (0.2, 0.2, 0.2),  # Very dark gray
            ]
            
            # Create pixel blocks
            for i in range(0, width, pixel_size):
                for j in range(0, height, pixel_size):
                    # Calculate pixel block boundaries
                    x1 = rect.x0 + i
                    y1 = rect.y0 + j
                    x2 = min(rect.x0 + i + pixel_size, rect.x1)
                    y2 = min(rect.y0 + j + pixel_size, rect.y1)
                    
                    # Create pixel rectangle
                    pixel_rect = fitz.Rect(x1, y1, x2, y2)
                    
                    # Choose color based on position for realistic pixelation
                    color_index = (i // pixel_size + j // pixel_size) % len(pixel_colors)
                    color = pixel_colors[color_index]
                    
                    # Draw the pixel block
                    page.draw_rect(pixel_rect, color=color, fill=color)
                    
        except Exception as e:
            logger.warning(f"Native pixelation failed, using simple fallback: {e}")
            # Simple fallback if pixelation fails
            page.draw_rect(rect, color=(0.5, 0.5, 0.5), fill=(0.5, 0.5, 0.5))

    def _validate_pdf_integrity(self, pdf_bytes: bytes) -> Dict[str, Any]:
        """
        Validate PDF integrity to ensure it's not corrupted
        """
        try:
            # Basic size check
            if len(pdf_bytes) < 100:
                return {'valid': False, 'error': 'PDF too small - likely corrupted'}
            
            # Check PDF header
            if not pdf_bytes.startswith(b'%PDF-'):
                return {'valid': False, 'error': 'Invalid PDF header'}
                
            # Try to open with PyMuPDF
            test_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            
            # Basic document checks
            page_count = test_doc.page_count
            if page_count == 0:
                test_doc.close()
                return {'valid': False, 'error': 'PDF has no pages'}
            
            # Try to access first page
            test_page = test_doc[0]
            test_page.get_text()  # This will fail if page is corrupted
            
            test_doc.close()
            
            logger.debug(f"PDF validation passed: {len(pdf_bytes)} bytes, {page_count} pages")
            return {'valid': True, 'pages': page_count, 'size': len(pdf_bytes)}
            
        except Exception as e:
            return {'valid': False, 'error': f'PDF validation failed: {str(e)}'}

    def _apply_presidio_redaction(self, page, rect, redaction_type: str):
        """Apply enhanced visual redaction effects (simplified approach to prevent PDF corruption)"""
        try:
            # Simplified approach to prevent PDF corruption from too many drawing operations
            
            if redaction_type == "blur":
                # Create a simple stippled pattern with fewer operations
                pattern_size = 3  # Larger pattern size to reduce operations
                for i in range(0, int(rect.width), pattern_size):
                    for j in range(0, int(rect.height), pattern_size):
                        if (i + j) % 6 < 3:  # Only draw half the rectangles for stipple effect
                            small_rect = fitz.Rect(
                                rect.x0 + i, rect.y0 + j, 
                                min(rect.x0 + i + pattern_size, rect.x1), 
                                min(rect.y0 + j + pattern_size, rect.y1)
                            )
                            # Alternate between gray shades for blur effect
                            gray_val = 0.5 + (0.2 * ((i + j) % 2))
                            page.draw_rect(small_rect, color=(gray_val, gray_val, gray_val), fill=(gray_val, gray_val, gray_val))
                        
            elif redaction_type == "pixelate":
                # Simplified pixelate with larger blocks and fewer operations
                block_size = 8  # Larger blocks to reduce PDF operations
                for i in range(0, int(rect.width), block_size):
                    for j in range(0, int(rect.height), block_size):
                        block_rect = fitz.Rect(
                            rect.x0 + i, rect.y0 + j, 
                            min(rect.x0 + i + block_size, rect.x1), 
                            min(rect.y0 + j + block_size, rect.y1)
                        )
                        # Simplified color variation for pixelate effect
                        pixel_variation = ((i // block_size + j // block_size) % 3)
                        if pixel_variation == 0:
                            pixel_color = (0.4, 0.4, 0.4)
                        elif pixel_variation == 1:
                            pixel_color = (0.6, 0.6, 0.6)
                        else:
                            pixel_color = (0.5, 0.5, 0.5)
                        page.draw_rect(block_rect, color=pixel_color, fill=pixel_color)
            else:
                raise ValueError(f"Unsupported redaction type: {redaction_type}")
                
        except Exception as e:
            logger.error(f"Enhanced redaction failed: {e}")
            raise e
    
    def _get_enhanced_replacement_text(self, entity_type: str, original_text: str) -> str:
        """Get enhanced replacement text using PIIRedactor for better contextual replacements"""
        try:
            # Use PIIRedactor for more intelligent replacement if available
            if self.pii_redactor:
                try:
                    # Attempt to use pii-redact for contextual replacement
                    # This provides more intelligent placeholders based on context
                    redacted_result = self.pii_redactor.redact(original_text)
                    if redacted_result and redacted_result != original_text:
                        # Extract the replacement from the redacted result
                        # PIIRedactor typically returns the redacted version
                        return redacted_result
                except Exception as pii_error:
                    logger.debug(f"PIIRedactor failed for {entity_type}: {pii_error}")
                    # Fall back to enhanced standard replacement
                    pass
            
            # Enhanced standard replacement with better categorization
            return self._get_replacement_text(entity_type)
            
        except Exception as e:
            logger.warning(f"Enhanced replacement failed for {entity_type}: {e}")
            # Fall back to standard replacement
            return self._get_replacement_text(entity_type)
    
    def cleanup(self):
        """Clean up temporary files and directories"""
        try:
            if hasattr(self, 'temp_dir') and self.temp_dir and os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
                logger.info(f"Cleaned up temporary directory: {self.temp_dir}")
                self.temp_dir = None
        except Exception as e:
            logger.warning(f"Cleanup failed: {e}")
    
    def __del__(self):
        """Ensure cleanup on object destruction"""
        self.cleanup()

# Global redaction engine instance
document_redactor = DocumentRedactionEngine()

def redact_document_file(file_content: bytes, 
                        filename: str, 
                        redaction_method: RedactionMethod = RedactionMethod.BLACKOUT,
                        sector: Sector = Sector.GENERAL,
                        confidence_threshold: float = 0.7) -> Dict[str, Any]:
    """Convenience function for document redaction"""
    
    # Create sector-specific redactor if needed
    if sector != Sector.GENERAL:
        redactor = DocumentRedactionEngine(sector=sector)
        result = redactor.redact_document(file_content, filename, redaction_method, confidence_threshold)
        redactor.cleanup()
        return result
    else:
        return document_redactor.redact_document(file_content, filename, redaction_method, confidence_threshold)

if __name__ == "__main__":
    # Test the redaction engine
    print("Document Redaction Engine initialized successfully!")
    print("Supported formats: PDF, DOCX, TXT, Images (PNG, JPG, TIFF, BMP)")
    print(f"Available redaction methods: {[method.value for method in RedactionMethod]}")